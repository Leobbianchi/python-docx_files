from datetime import datetime
from docx import Document
from num2words import num2words

dict_valor_parcelas = {
    '1' : 'uma',
    '2' : 'duas',
    '3' : 'três',
    '4' : 'quatro',
    '5' : 'cinco'
}

def ajuste_str(valor):
    """
    Function used to format the values into brazilian-real currency.
    """
    if len(valor) == 10:
        resultado = valor[:4] + '.' + valor[4:]
    elif len(valor) == 11:
        resultado = valor[:5] + '.' + valor[5:]
    else:
        resultado = valor
    return resultado

def number_to_long_number(number_p):
    """
    Function used to write the numbers.
    """
    if number_p.find(',') != -1:
        number_p = number_p.split(',')
        number_p1 = int(number_p[0].replace('.', ''))
        number_p2 = int(number_p[1])
    else:
        number_p1 = int(number_p.replace('.', ''))
        number_p2 = 0

    if number_p1 == 1:
        aux1 = ' real'
    else:
        aux1 = ' reais'

    if number_p2 == 1:
        aux2 = ' centavo'
    else:
        aux2 = ' centavos'

    text1 = ''
    if number_p1 > 0:
        text1 = num2words(number_p1, lang='pt_BR') + str(aux1)
    else:
        text1 = ''

    if number_p2 > 0:
        text2 = num2words(number_p2, lang='pt_BR') + str(aux2)
    else:
        text2 = ''

    if (number_p1 > 0 and number_p2 > 0):
        result = text1 + ' e ' + text2
    else:
        result = text1 + text2

    result_str = str(result).replace(',','')
    split = result_str.split(' ')

    if split[0] == 'mil':
        split[0] = 'um mil'

    resultado_final = ' '.join(split)

    return resultado_final

# Dictionary used to transform the months
dic_mes = {
            1 : 'janeiro',
            2 : 'fevereiro',
            3 : 'março',
            4 : 'abril',
            5 : 'maio',
            6 : 'junho',
            7 : 'julho',
            8 : 'agosto',
            9 : 'setembro',
            10 : 'outubro',
            11 : 'novembro',
            12 : 'dezembro'
}


# ask for inputs:
validacao = 'N'

while validacao == 'N':
    # NOME
    nome = input('Qual o nome da cliente? >').title()
    # CPF
    cpf = input('Qual o CPF? (apenas numeros) >')

    while len(cpf) != 11: # checando se o cpf foi digitado corretamente
        cpf = input('CPF invalido, digite um CPF de 11 digitos (apenas numeros) >')
    cpf = cpf[0:3] + '.' + cpf[3:6] + '.' + cpf[6:9] + '-' + cpf[9:11]

    # ENDERECO
    endereco = input('Qual o endereço da cliente >')

    # VALOR TOTAL
    valor = input('Qual o valor total do vestido? (apenas numeros) >')
    valor_str2 = 'R$ ' + valor + ',00'
    valor_str = ajuste_str(valor_str2)

    valor_extenso = number_to_long_number(valor)

    valor_desconto_a_vista = round(int(valor) * 0.96,0)
    valor_desconto_a_vista_str2 = 'R$ ' + str(valor_desconto_a_vista) + '0'
    valor_desconto_a_vista_str3 = valor_desconto_a_vista_str2.replace('.',',')
    valor_desconto_a_vista_str = ajuste_str(valor_desconto_a_vista_str3)

    # QUANTIDADE DE PARCELAS
    parcelas = input('Quantas parcelas? >')
    parcelas_extenso = dict_valor_parcelas[parcelas]

    # DATA DO CASAMENTO                 )
    data_casamento = input('Qual a data do casamento ? DD/MM/AAAA >')
    dia_casamento, mes_casamento, ano_casamento = data_casamento.split('/')
    dia_casamento = int(dia_casamento)
    mes_casamento = int(mes_casamento)
    ano_casamento = int(ano_casamento)

    # DIA ATUAL
    data = datetime.now().strftime("%d/%m/%Y")

    # CONFIRMAR OS DADOS - COMPLETAR
    validacao = input(f'Verifique se os valores estao corretos:\nNome: {nome}\nCPF: {cpf}\nValor: {valor}\nS ou N? > ').title()

# VALOR DA PARCELA
valor_parcela = int(valor) / int(parcelas)
valor_parcela_str2 = 'R$ ' + str(valor_parcela) +'0'
valor_parcela_str3 = valor_parcela_str2.replace('.',',')
valor_parcela_str = ajuste_str(valor_parcela_str3)
valor_parcela2 = str(valor_parcela).replace('.',',')
valor_parcelas_extenso = number_to_long_number(str(valor_parcela2))

# VALOR DA MULTA
valor_multa = round(int(valor) * 0.2,0)
valor_multa_str2 = 'R$ '+ str(valor_multa) + '0'
valor_multa_str3 = valor_multa_str2.replace('.',',')
valor_multa_str = ajuste_str(valor_multa_str3)

# DATA DA ENTREGA
if dia_casamento > 6:
    dia_entrega = dia_casamento - 5
    mes_entrega = mes_casamento
    ano_entrega = ano_casamento
    data_entrega = str(dia_entrega) + ' de ' + dic_mes[mes_entrega] + ' de ' + str(ano_entrega)
else:
    dia_entrega = 30 + (dia_casamento - 5)
    if mes_casamento == 1:
        mes_entrega = 12
        ano_entrega = ano_casamento - 1
    else:
        mes_entrega = mes_casamento - 1
        ano_entrega = ano_casamento
    data_entrega = str(dia_entrega) + ' de ' + dic_mes[mes_entrega] + ' de ' + str(ano_entrega)

# AJUSTAR PARCELAS
if int(parcelas) == 1:

    # open word file
    document = Document('vista.docx')
    # VALOR DA ENTRADA
    valor_entrada = round(int(valor_desconto_a_vista) * 0.35, 0)
    validacao = input(f'O valor da entrada ficou em R$ {valor_entrada} que representa exatamente 35% do valor total.\nCaso queira usar outro valor, digite apenas numeros, caso contrario digite ok >')
    if validacao != 'ok':
        valor_entrada = validacao

    valor_entrada_str = 'R$ ',str(valor_entrada),',00'
    valor_entrada_str = ''.join(valor_entrada_str)

    # DATA DA ENTRADA
    if mes_entrega < 4:
        mes_entrada = 12 + (mes_entrega - 3)
        ano_entrada = ano_entrega
    else:
        ano_entrada = ano_entrega
        mes_entrada = mes_entrega - 3

    data_entrada = '15/' + str(mes_entrada) + '/' + str(ano_entrada)

    # VALOR DO SALDO (TOTAL - ENTRADA)
    valor_saldo = int(valor_desconto_a_vista) - int(valor_entrada)
    valor_saldo_str = 'R$ ',str(int(valor_desconto_a_vista) - int(valor_entrada)),',00'
    valor_saldo_str = ''.join(valor_saldo_str)
    # DATA DO PAGAMENTO DO SALDO
    data_saldo = data_entrega
    dict_trocas = {
            'nomecliente': str(nome),
            'cpaqui': str(cpf),
            'endcliente': str(endereco),
            'dataentrega': data_entrega,
            'valortot': str(valor_str),
            'totextenso': valor_extenso,
            'valdescon': str(valor_desconto_a_vista_str),
            'valentrada': str(valor_entrada_str),
            'mesientrada': str(dic_mes[mes_entrada]),
            'anoientrada': str(ano_entrada),
            '[saldo]': str(valor_saldo_str),
            'mesfparcela': str(dic_mes[mes_entrega]),
            'anofparcela': str(ano_entrega),
            '[multa]': str(valor_multa_str),
            '[data]': str(data),
        }
    chaves_parc = 0
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            for key in dict_trocas.keys():
                if key in text:
                    chaves_parc += 1
                    text = text.replace(key, dict_trocas[key])
                    inline[i].text = text
    print('Foram inseridos ', chaves_parc, 'dados, de um total de 15')
else:
        # open word file
    document = Document("parcelado.docx")
    comparativo = (int(mes_entrega) - int(parcelas))
    if comparativo >= 0:
        mes_inicio_parcela = comparativo + 1
        ano_inicio_parcela = ano_entrega

    else:
        mes_inicio_parcela = 12 + comparativo + 1
        ano_inicio_parcela = ano_entrega - 1
    data_inicio_parcela = '15/' + str(mes_inicio_parcela) + '/' + str(ano_inicio_parcela)
    dict_trocas = {
            'nomecliente': str(nome),
            'cpaqui': str(cpf),
            'endcliente': str(endereco),
            'dataentrega': data_entrega,
            'valortot': str(valor_str),
            'totextenso': valor_extenso,
            'parcextens' : str(parcelas_extenso),
            'valparc' : str(valor_parcela_str),
            '[parcelas]' : str(parcelas),
            'valpext': str(valor_parcelas_extenso),
            'mesprimpar': dic_mes[mes_inicio_parcela],
            'anoprimpar': str(ano_inicio_parcela),
            'mesentr': str(dic_mes[mes_entrega]),
            'anoentr': str(ano_entrega),
            '[multa]': str(valor_multa_str),
            '[data]': str(data),
        }
    chaves_parc = 0
    for p in document.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            for key in dict_trocas.keys():
                if key in text:
                    chaves_parc += 1
                    text = text.replace(key, dict_trocas[key])
                    inline[i].text = text
    print('Foram inseridos ', chaves_parc, 'dados, de um total de 16')




# Exportar o arquivo com o primeiron nome _contrato.docx
nome_contrato_completo = nome.split()
nome_contrato = nome_contrato_completo[0]
nome_arquivo = f'{nome_contrato}_contrato.docx'
document.save(nome_arquivo)
